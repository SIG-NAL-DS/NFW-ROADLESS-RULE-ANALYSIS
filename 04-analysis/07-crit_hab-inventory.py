"""
Roadless Areas – Critical Habitat (Simple Reporting Only)
Layer: crithab_poly_ra (already intersected with roadless areas)

Outputs (no modeling / no ranking beyond simple counts):
- CSV + LaTeX for each table
- Single Word doc containing all tables
"""

# ---------------------- IMPORTS ----------------------
from pathlib import Path
import pandas as pd
import geopandas as gpd
from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches

from nfw_project import config_paths as cfg


# ---------------------- CONFIG ----------------------
OUTPUT_DIR = Path(cfg.TABLES_OUTPUT_DIR) / "roadless_critical_habitat_tables"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

CRITHAB_LAYER = "crithab_poly_ra"
USFS_ADMIN_RA_LAYER = "usfs_admin_boundary_RA"
STATES_RA_LAYER = "States_RA"


# ------------------ EXPORT HELPERS ------------------
def export_df_to_all_formats(df: pd.DataFrame, base_name: str, output_dir: Path) -> None:
    """
    Export a DataFrame to CSV and LaTeX.
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    df.to_csv(output_dir / f"{base_name}.csv", index=False)

    tex_str = df.to_latex(index=False, escape=True)
    (output_dir / f"{base_name}.tex").write_text(tex_str, encoding="utf-8")


def add_df_to_word_doc(doc: DocxDocument, df: pd.DataFrame, title: str) -> None:
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    # header
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)

    # rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            cells[i].text = "" if pd.isna(val) else str(val)

    doc.add_paragraph()
    doc.add_paragraph()

def set_doc_margins_half_inch(doc: DocxDocument) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

def export_all_tables_to_word(tables: dict[str, pd.DataFrame], output_path: Path) -> None:
    doc = Document()
    set_doc_margins_half_inch(doc)
    doc.add_heading("Roadless Areas – Critical Habitat", level=1)

    for title, df in tables.items():
        # Start Species Inventory on a new page
        if title.startswith("Table G."):
            doc.add_page_break()

        add_df_to_word_doc(doc, df, title)

    doc.save(str(output_path))



# ------------------ DATA LOADING ------------------
def load_data():
    crit_hab = gpd.read_file(cfg.ROADLESS_ANALYSIS_GPKG, layer=CRITHAB_LAYER)
    usfs_admin = gpd.read_file(cfg.ROADLESS_ANALYSIS_GPKG, layer=USFS_ADMIN_RA_LAYER)
    states = gpd.read_file(cfg.ROADLESS_ANALYSIS_GPKG, layer=STATES_RA_LAYER)

    # Force everything to EPSG:5070 for consistent overlay/join behavior
    if crit_hab.crs is None or usfs_admin.crs is None or states.crs is None:
        raise ValueError("One or more layers are missing CRS information.")

    if crit_hab.crs.to_string() != "EPSG:5070":
        crit_hab = crit_hab.to_crs("EPSG:5070")
    if usfs_admin.crs.to_string() != "EPSG:5070":
        usfs_admin = usfs_admin.to_crs("EPSG:5070")
    if states.crs.to_string() != "EPSG:5070":
        states = states.to_crs("EPSG:5070")


    return crit_hab, usfs_admin, states



# ------------------ REPORTING TABLES ------------------
def build_table_a_national_summary(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    total_polys = len(gdf)
    unique_sciname = gdf["sciname"].nunique() if "sciname" in gdf.columns else None
    unique_comname = gdf["comname"].nunique() if "comname" in gdf.columns else None
    unique_status = gdf["status"].nunique() if "status" in gdf.columns else None

    # Multi-species designations (best-effort: contains "multi")
    if "singlmulti" in gdf.columns:
        sm = gdf["singlmulti"].astype(str).str.lower()
        multi_polys = int(sm.str.contains("multi", na=False).sum())
    else:
        multi_polys = None

    rows = [
        ("Total critical habitat polygons intersecting roadless areas", total_polys),
        ("Unique species (scientific name)", unique_sciname),
        ("Unique species (common name)", unique_comname),
        ("Unique ESA status values", unique_status),
        ("Polygons flagged as multi-species designations (singlmulti contains 'multi')", multi_polys),
    ]
    return pd.DataFrame(rows, columns=["Metric", "Value"])


def build_table_b_species_inventory(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    cols = [c for c in ["sciname", "comname", "status", "listing_st", "spcode", "vipcode", "entity_id"] if c in gdf.columns]
    group_cols = [c for c in ["sciname", "comname", "status"] if c in gdf.columns]

    # Fallback grouping if missing
    if not group_cols:
        group_cols = [c for c in ["entity_id", "spcode", "vipcode"] if c in gdf.columns]
    if not group_cols:
        raise ValueError("Cannot build species inventory: no species identifier columns found.")

    agg = {}
    if "unit" in gdf.columns:
        agg["Units"] = ("unit", "nunique")
    if "subunit" in gdf.columns:
        agg["Subunits"] = ("subunit", "nunique")
    agg["Polygons"] = ("geometry", "count")

    df = (
        gdf
        .groupby(group_cols, dropna=False)
        .agg(**agg)
        .reset_index()
        .sort_values(["Polygons"], ascending=False)
    )

    # Add a stable “Species” label if possible
    if "sciname" in df.columns and "comname" in df.columns:
        df.insert(0, "Species", df["comname"].fillna("").astype(str) + " (" + df["sciname"].fillna("").astype(str) + ")")

    # Keep table compact and publication-friendly
    keep = [c for c in ["Species", "sciname", "comname", "status", "listing_st", "Units", "Subunits", "Polygons"] if c in df.columns]
    return df[keep]

def build_table_c_by_status(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    if "status" not in gdf.columns:
        return pd.DataFrame([("status field not present",)], columns=["Note"])

    df = (
        gdf
        .groupby("status", dropna=False)
        .agg(
            Species=("sciname", "nunique") if "sciname" in gdf.columns else ("geometry", "count"),
            Units=("unit", "nunique") if "unit" in gdf.columns else ("geometry", "count"),
            Polygons=("geometry", "count"),
        )
        .reset_index()
        .sort_values("Polygons", ascending=False)
        .rename(columns={"status": "ESA Status"})
    )
    return df

def build_table_c_by_listing_status(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    if "listing_st" not in gdf.columns:
        return pd.DataFrame([("listing_st field not present",)], columns=["Note"])

    df = (
        gdf.groupby("listing_st", dropna=False)
        .size()
        .reset_index(name="Polygons")
        .sort_values("Polygons", ascending=False)
        .rename(columns={"listing_st": "Listing Status"})
    )
    return df


def build_table_d_critical_hab_by_region(
    crit_hab: gpd.GeoDataFrame,
    usfs_admin: gpd.GeoDataFrame
) -> pd.DataFrame:
    if "REGION" not in usfs_admin.columns:
        return pd.DataFrame([("REGION field not present in usfs_admin_boundary_RA",)], columns=["Note"])

    ch = crit_hab.copy()
    ch["CH_ID"] = ch.index.astype(int)

    joined = gpd.sjoin(
        ch[["CH_ID", "sciname", "geometry"]],
        usfs_admin[["REGION", "geometry"]],
        how="left",
        predicate="intersects"
    )

    # Diagnostics
    unmatched = int(joined["REGION"].isna().sum())
    if unmatched > 0:
        print(f"Warning: {unmatched} critical habitat polygons did not match a USFS REGION.")

    out = (
        joined.groupby("REGION", dropna=False)
        .agg(
            **{
                "Critical Habitat Polygons": ("CH_ID", "nunique"),
                "Unique Species": ("sciname", "nunique"),
            }
        )
        .reset_index()
        .rename(columns={"REGION": "USFS Region"})
        .sort_values("Critical Habitat Polygons", ascending=False)
    )

    return out



def build_table_e_forest_count(usfs_admin: gpd.GeoDataFrame) -> pd.DataFrame:
    if "FORESTNUMB" in usfs_admin.columns:
        n_forests = usfs_admin["FORESTNUMB"].nunique()
    elif "FORESTNAME" in usfs_admin.columns:
        n_forests = usfs_admin["FORESTNAME"].nunique()
    else:
        n_forests = len(usfs_admin)

    return pd.DataFrame(
        [("Number of unique forests with critical habitat (in RA)", n_forests)],
        columns=["Metric", "Value"]
    )



def build_table_f_critical_hab_by_state(
    crit_hab: gpd.GeoDataFrame,
    states: gpd.GeoDataFrame
) -> pd.DataFrame:
    state_col = None
    for cand in ["STUSPS", "STATE", "STATE_ABBR", "NAME"]:
        if cand in states.columns:
            state_col = cand
            break
    if state_col is None:
        return pd.DataFrame([("No recognized state field found in States_RA",)], columns=["Note"])

    ch = crit_hab.copy()
    ch["CH_ID"] = ch.index.astype(int)

    joined = gpd.sjoin(
        ch[["CH_ID", "sciname", "geometry"]],
        states[[state_col, "geometry"]],
        how="left",
        predicate="intersects"
    )

    unmatched = int(joined[state_col].isna().sum())
    if unmatched > 0:
        print(f"Warning: {unmatched} critical habitat polygons did not match a state in States_RA.")

    out = (
        joined.groupby(state_col, dropna=False)
        .agg(
            **{
                "Critical Habitat Polygons": ("CH_ID", "nunique"),
                "Unique Species": ("sciname", "nunique"),
            }
        )
        .reset_index()
        .rename(columns={state_col: "State"})
        .sort_values("Critical Habitat Polygons", ascending=False)
    )

    return out


    out = (
        joined.groupby(state_col, dropna=False)
        .agg(
            **{
                "Critical Habitat Polygons": ("CH_ID", "nunique"),
                "Unique Species": ("sciname", "nunique"),
            }
        )
        .reset_index()
        .rename(columns={state_col: "State"})
        .sort_values("Critical Habitat Polygons", ascending=False)
    )

    return out


def build_table_g_species_inventory_with_state(
    crit_hab: gpd.GeoDataFrame,
    states: gpd.GeoDataFrame
) -> pd.DataFrame:
    # Find a state column on the states layer
    state_col = None
    for cand in ["STUSPS", "STATE", "STATE_ABBR", "NAME"]:
        if cand in states.columns:
            state_col = cand
            break
    if state_col is None:
        return pd.DataFrame([("No recognized state field found in States_RA",)], columns=["Note"])

    # Join states onto critical habitat polygons
    keep = [c for c in ["comname", "sciname", "status", "listing_st", "geometry"] if c in crit_hab.columns]
    ch = crit_hab[keep].copy()

    joined = gpd.sjoin(
        ch,
        states[[state_col, "geometry"]],
        how="left",
        predicate="intersects"
    ).rename(columns={state_col: "State"})

    # Keep a clean, long-form inventory (distinct combinations)
    cols = [c for c in ["State", "listing_st", "status", "comname", "sciname"] if c in joined.columns]

    out = (
        joined[cols]
        .drop_duplicates()
        .sort_values([c for c in ["State", "status", "sciname"] if c in cols])
        .reset_index(drop=True)
    )

    return out



# ------------------------ MAIN ----------------------
def main():
    print("Loading datasets...")
    crit_hab, usfs_admin, states = load_data()

    print("Building reporting tables...")
    tA = build_table_a_national_summary(crit_hab)

    # 2nd table: ESA status summary
    tB = build_table_c_by_status(crit_hab)

    # counts by listing_st
    tC = build_table_c_by_listing_status(crit_hab)

    # admin boundary counts by REGION
    tD = build_table_d_critical_hab_by_region(crit_hab, usfs_admin)


    # forest count summary (single row)
    tE = build_table_e_forest_count(usfs_admin)

    # critical habitats per state + unique species per state (spatial join only)
    tF = build_table_f_critical_hab_by_state(crit_hab, states)

    # species inventory LAST (no units/subunits/polygons; includes listing_st + state)
    tG = build_table_g_species_inventory_with_state(crit_hab, states)

    tables = {
        "Table A. National Summary (Critical Habitat ∩ Roadless Areas)": tA,
        "Table B. Summary by ESA Status": tB,
        "Table C. Counts by Listing Status (listing_st)": tC,
        "Table D. Roadless Admin Boundary Counts by USFS Region": tD,
        "Table E. Number of Unique Forests with Critical Habitat (in RA)": tE,
        "Table F. Critical Habitat Polygons and Unique Species by State": tF,
        "Table G. Species Inventory (Critical Habitat overlapping Roadless Areas)": tG,
    }

    print("Exporting tables to CSV / LaTeX...")
    export_df_to_all_formats(tA, "tableA_national_summary", OUTPUT_DIR)
    export_df_to_all_formats(tB, "tableB_by_esa_status", OUTPUT_DIR)
    export_df_to_all_formats(tC, "tableC_by_listing_status", OUTPUT_DIR)
    export_df_to_all_formats(tD, "tableD_usfs_admin_by_region", OUTPUT_DIR)
    export_df_to_all_formats(tE, "tableE_unique_forest_count", OUTPUT_DIR)
    export_df_to_all_formats(tF, "tableF_critical_hab_by_state", OUTPUT_DIR)
    export_df_to_all_formats(tG, "tableG_species_inventory", OUTPUT_DIR)

    print("Exporting all tables to Word...")
    export_all_tables_to_word(
        tables,
        OUTPUT_DIR / "roadless_critical_habitat_reporting_tables.docx"
    )

    print("Done.")
    print(f"Outputs written to: {OUTPUT_DIR.resolve()}")




if __name__ == "__main__":
    main()
