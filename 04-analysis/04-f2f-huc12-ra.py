"""
Roadless Areas – F2F HUC12 Model Output (Category E: Risk Indices)
Layer: F2F2_HUC12_RA (from cfg.ROADLESS_ANALYSIS_GPKG)

Goal (this branch):
- Aggregate HUC12-modeled "risk/index" fields (Category E) to Roadless Areas.
- Prefer area-weighted MEAN when overlap fractions can be computed reliably.
- If weighting cannot be computed, report RAW (unweighted) summary stats instead.
- Output Top-20 tables (CSV + LaTeX + Word).

No Markdown exports.
"""

# ---------------------- IMPORTS ----------------------
from pathlib import Path
import numpy as np
import pandas as pd
import geopandas as gpd
from docx import Document
from docx.document import Document as DocxDocument

from nfw_project import config_paths as cfg


# ---------------------- CONFIG ----------------------
OUTPUT_DIR = Path(cfg.TABLES_OUTPUT_DIR) / "roadless_f2f2_huc12_risk_tables"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

ROADLESS_LAYER = "roadless_area"
MODEL_LAYER = "F2F2_HUC12_RA"  # <-- new layer


# Category E fields (modeled indices / risk components)
# NOTE: We do not assume directionality (higher=better/worse). We will label as "Highest index value".
RISK_FIELDS = [
    "WFP", "IDRISK",
    "R_AG", "R_RIP", "R_IMPV", "R_NATCOV", "R_Q",
    "IMP", "IMP_R",
    "APCW", "APCW_R",
    "WFP_IMP_R", "IDRISK_R",
]

# We will NOT aggregate *_R as numeric sums; but we can still report raw mean/max if present.
# If you want to exclude all *_R from outputs, set EXCLUDE_RANK_FIELDS=True.
EXCLUDE_RANK_FIELDS = False


# ------------------ EXPORT HELPERS ------------------
def export_df_to_all_formats(df: pd.DataFrame, base_name: str, output_dir: Path) -> None:
    """
    Export a DataFrame to CSV and LaTeX.
    Outputs:
      - {base_name}.csv
      - {base_name}.tex
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    df.to_csv(output_dir / f"{base_name}.csv", index=False)

    tex_str = df.to_latex(index=False, escape=True)
    (output_dir / f"{base_name}.tex").write_text(tex_str, encoding="utf-8")


def add_df_to_word_doc(doc: DocxDocument, df: pd.DataFrame, title: str) -> None:
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    # header
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)

    # rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = "" if pd.isna(row[col]) else str(row[col])

    doc.add_paragraph()
    doc.add_paragraph()


def export_all_tables_to_word(tables: dict[str, pd.DataFrame], output_path: Path) -> None:
    doc = Document()
    doc.add_heading("Roadless Areas – F2F HUC12 Model (Risk Indices: Category E)", level=1)

    for title, df in tables.items():
        add_df_to_word_doc(doc, df, title)

    doc.save(str(output_path))


# ------------------ DATA LOADING ------------------
def load_data():
    roadless = gpd.read_file(cfg.USFS_GPKG, layer=ROADLESS_LAYER)
    model = gpd.read_file(cfg.ROADLESS_ANALYSIS_GPKG, layer=MODEL_LAYER)

    # Align CRS to EPSG:5070 for area-safe operations
    if roadless.crs is None:
        raise ValueError("roadless.crs is None; CRS must be defined.")
    if model.crs is None:
        raise ValueError("model.crs is None; CRS must be defined.")

    if roadless.crs.to_string() != "EPSG:5070":
        roadless = roadless.to_crs("EPSG:5070")

    if model.crs.to_string() != "EPSG:5070":
        model = model.to_crs("EPSG:5070")

    # Stable Roadless Area ID
    if "RA_ID" not in roadless.columns:
        roadless = roadless.reset_index(drop=True)
        roadless["RA_ID"] = roadless.index + 1

    return roadless, model



def coerce_numeric(df: pd.DataFrame, cols: list[str]) -> pd.DataFrame:
    for c in cols:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")
    return df


# ------------------ ANALYSIS ------------------
def align_to_epsg5070(roadless: gpd.GeoDataFrame, model: gpd.GeoDataFrame):
    if roadless.crs is None:
        raise ValueError("roadless.crs is None; set CRS before overlay.")
    if model.crs is None:
        raise ValueError("model.crs is None; set CRS before overlay.")

    if str(roadless.crs).lower() != "epsg:5070":
        roadless = roadless.to_crs("EPSG:5070")
    if str(model.crs).lower() != "epsg:5070":
        model = model.to_crs("EPSG:5070")

    return roadless, model


def build_ra_huc12_overlay(roadless: gpd.GeoDataFrame, model: gpd.GeoDataFrame) -> gpd.GeoDataFrame:
    """
    Overlay roadless polygons with HUC12 model polygons so we can:
    - avoid double counting where an RA touches multiple HUC12s
    - compute overlap fractions for optional weighted means

    This is robust even if MODEL_LAYER was already prefiltered by intersection.
    """
    keep_ra = ["RA_ID", "REGION", "FOREST", "STATE", "NAME", "geometry"]
    keep_ra = [c for c in keep_ra if c in roadless.columns]

    # Determine area column in model. Prefer 'Acres'. If missing, we can still do raw stats.
    keep_model = ["HUC12", "Acres", "geometry"] + [c for c in RISK_FIELDS if c in model.columns]
    keep_model = list(dict.fromkeys([c for c in keep_model if c in model.columns] + ["geometry"]))

    ra_huc12 = gpd.overlay(
        roadless[keep_ra],
        model[keep_model],
        how="intersection"
    )

    # Intersection area acres
    ra_huc12["INT_ACRES"] = ra_huc12.geometry.area / 4046.8564224

    # Overlap fraction of HUC12 (only if model has Acres and it's usable)
    if "Acres" in ra_huc12.columns:
        ra_huc12["FRAC_HUC_IN_RA"] = ra_huc12["INT_ACRES"] / ra_huc12["Acres"]
        ra_huc12["FRAC_HUC_IN_RA"] = ra_huc12["FRAC_HUC_IN_RA"].clip(lower=0, upper=1)
    else:
        ra_huc12["FRAC_HUC_IN_RA"] = np.nan

    return ra_huc12


def summarize_category_e_by_ra(ra_huc12: gpd.GeoDataFrame) -> pd.DataFrame:
    """
    For Category E indices:
    - Preferred: area-weighted MEAN using FRAC_HUC_IN_RA (if computable)
    - Always: RAW (unweighted) MEAN and MAX across intersecting HUC12 pieces
    """
    # Decide which risk fields we actually use
    fields = [c for c in RISK_FIELDS if c in ra_huc12.columns]
    if EXCLUDE_RANK_FIELDS:
        fields = [c for c in fields if not c.endswith("_R") and "_R" not in c]

    # Base grouping keys
    keys = ["RA_ID", "REGION", "FOREST", "STATE", "NAME"]

    # Always compute RA acres overlapped + HUC12 count
    out = (
        ra_huc12
        .groupby(keys, dropna=False)
        .agg(
            RA_Acres_in_Model=("INT_ACRES", "sum"),
            HUC12_Count=("HUC12", "nunique") if "HUC12" in ra_huc12.columns else ("INT_ACRES", "count"),
        )
        .reset_index()
    )

    # RAW mean + max (always safe)
    raw_mean = ra_huc12.groupby(keys, dropna=False)[fields].mean(numeric_only=True).add_suffix("_raw_mean").reset_index()
    raw_max  = ra_huc12.groupby(keys, dropna=False)[fields].max(numeric_only=True).add_suffix("_raw_max").reset_index()

    out = out.merge(raw_mean, on=keys, how="left").merge(raw_max, on=keys, how="left")

    # Optional: area-weighted mean (only if FRAC_HUC_IN_RA is present and not all missing)
    can_weight = ("FRAC_HUC_IN_RA" in ra_huc12.columns) and ra_huc12["FRAC_HUC_IN_RA"].notna().any()

    if can_weight:
        # Weighted mean = sum(value * frac) / sum(frac)
        denom = ra_huc12.groupby(keys, dropna=False)["FRAC_HUC_IN_RA"].sum().replace(0, np.nan)

        wparts = {}
        for f in fields:
            wparts[f] = (ra_huc12[f] * ra_huc12["FRAC_HUC_IN_RA"])

        wdf = pd.DataFrame(wparts)
        wdf[keys] = ra_huc12[keys].reset_index(drop=True)

        num = wdf.groupby(keys, dropna=False)[fields].sum(numeric_only=True)

        wmean = (num.div(denom, axis=0)).add_suffix("_aw_mean").reset_index()
        out = out.merge(wmean, on=keys, how="left")
    else:
        # If weighting doesn't work: we simply do not add aw columns.
        pass

    return out


def top20_by_metric(ra_summary: pd.DataFrame, metric_col: str, label_cols: list[str]) -> pd.DataFrame:
    """
    Build a Top 20 table by the *highest* values of metric_col.
    We do not assume directionality beyond "highest index value".
    """
    if metric_col not in ra_summary.columns:
        raise ValueError(f"Metric column not found: {metric_col}")

    df = (
        ra_summary
        .sort_values(metric_col, ascending=False)
        .head(20)
        .assign(Rank=lambda d: range(1, len(d) + 1))
        [["Rank"] + label_cols + [metric_col]]
    )

    # nice rounding (indices usually 0-1 or similar; keep 4 decimals)
    df[metric_col] = pd.to_numeric(df[metric_col], errors="coerce").round(4)

    return df


# ------------------------ MAIN ----------------------
def main():
    print("Loading datasets...")
    roadless, model = load_data()

    print("Coercing numeric fields (model)...")
    # Coerce any risk fields present + Acres
    model = coerce_numeric(model, ["Acres"] + [c for c in RISK_FIELDS if c in model.columns])


    print("Building RA–HUC12 overlay...")
    ra_huc12 = build_ra_huc12_overlay(roadless, model)

    print("Summarizing Category E risk indices by roadless area...")
    ra_risk = summarize_category_e_by_ra(ra_huc12)

    # Choose the metric variant for Top-20 tables:
    # Prefer area-weighted mean if present; otherwise raw mean.
    # (This respects your request: if weighting doesn't work, report raw.)
    def pick_metric(base_name: str) -> str:
        aw = f"{base_name}_aw_mean"
        raw = f"{base_name}_raw_mean"
        if aw in ra_risk.columns and ra_risk[aw].notna().any():
            return aw
        return raw

    # Labels to show in top tables
    label_cols = ["NAME", "FOREST", "STATE", "REGION", "RA_Acres_in_Model", "HUC12_Count"]

    # Build Top-20 tables for key Category E metrics (edit this list freely)
    targets = [
        ("WFP",   "Top 20 Roadless Areas by Highest WFP Index (Category E)"),
        ("IDRISK","Top 20 Roadless Areas by Highest IDRISK Index (Category E)"),
        ("R_Q",   "Top 20 Roadless Areas by Highest R_Q Index (Category E)"),
        ("R_IMPV","Top 20 Roadless Areas by Highest R_IMPV Index (Category E)"),
        ("R_AG",  "Top 20 Roadless Areas by Highest R_AG Index (Category E)"),
    ]

    top_tables = {}
    for base, title in targets:
        if base not in RISK_FIELDS:
            continue
        if base not in model.columns:
            print(f"Skipping {base}: not found in model layer.")
            continue

        metric = pick_metric(base)
        t = top20_by_metric(ra_risk, metric, label_cols)

        # Rename the metric column to something readable in Word/exports
        t = t.rename(columns={metric: f"{base} ({'area-weighted mean' if metric.endswith('_aw_mean') else 'raw mean'})"})
        top_tables[title] = t

        # Export each Top-20 to CSV/LaTeX
        safe_name = base.lower()
        export_df_to_all_formats(t, f"top20_{safe_name}_categoryE", OUTPUT_DIR)

    # OPTIONAL: export the full RA risk summary as CSV/LaTeX (can be large; keep off Word)
    # export_df_to_all_formats(ra_risk, "ra_categoryE_risk_summary_full", OUTPUT_DIR)

    print("Exporting all Top-20 Category E tables to Word...")
    export_all_tables_to_word(
        top_tables,
        OUTPUT_DIR / "roadless_f2f2_huc12_categoryE_top20_tables.docx"
    )

    print("Done.")
    print(f"Outputs written to: {OUTPUT_DIR.resolve()}")


if __name__ == "__main__":
    main()
