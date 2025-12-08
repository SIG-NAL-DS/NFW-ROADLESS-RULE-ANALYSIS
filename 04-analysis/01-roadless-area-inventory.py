"""
Roadless Area Inventory: Table Generation + Export

This script:
- Loads a roadless_5070 dataset
- Builds publication-ready tables (summary stats)
- Exports each table to CSV, Markdown, LaTeX
- Writes a single Word document with all tables

Adjust the CONFIG section as needed.
"""

# ---------------------- CONFIG ----------------------
from pathlib import Path
from nfw_project import config_paths as cfg

# Path to the GPKG that contains both layers
ROADLESS_PATH = cfg.USFS_GPKG

# Layer names inside the GPKG
ROADLESS_LAYER = "roadless_area"
FOREST_ADMIN_LAYER = "forest_admin_boundary"

# Output directory for tables
OUTPUT_DIR = cfg.TABLES_OUTPUT_DIR / "roadless_tables"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ------------------ IMPORTS & SETUP -----------------
import geopandas as gpd
import pandas as pd
import numpy as np

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches


# ----------------- LOADING DATA ---------------------
def load_roadless(path: Path, layer: str | None = None) -> gpd.GeoDataFrame:
    if layer:
        gdf = gpd.read_file(path, layer=layer)
    else:
        gdf = gpd.read_file(path)

    # Ensure required columns exist
    required = [
        "REGION", "FOREST", "STATE", "NAME",
        "CATEGORY", "ACRES", "SHAPE_AREA", "SHAPE_LEN", "geometry"
    ]
    missing_cols = [c for c in required if c not in gdf.columns]
    if missing_cols:
        raise ValueError(f"Missing required columns: {missing_cols}")

    # Make sure numeric fields are numeric
    for col in ["ACRES", "SHAPE_AREA", "SHAPE_LEN"]:
        gdf[col] = pd.to_numeric(gdf[col], errors="coerce")

    return gdf


# ----------------- TABLE BUILDERS -------------------
def build_table1_national_summary(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    total_polygons = len(gdf)
    total_acres = gdf["ACRES"].sum(skipna=True)

    desc_acres = gdf["ACRES"].describe()
    mean_acres = desc_acres["mean"]
    median_acres = desc_acres["50%"]
    min_acres = desc_acres["min"]
    max_acres = desc_acres["max"]

    n_forests = gdf["FOREST"].nunique()
    n_states = gdf["STATE"].nunique()
    n_regions = gdf["REGION"].nunique()

    data = [
        ("Total Number of Roadless Area Polygons", f"{total_polygons:,}"),
        ("Total Acreage", f"{total_acres:,.1f}"),
        ("Mean Polygon Size (acres)", f"{mean_acres:,.1f}"),
        ("Median Polygon Size (acres)", f"{median_acres:,.1f}"),
        ("Minimum Polygon Size (acres)", f"{min_acres:,.1f}"),
        ("Maximum Polygon Size (acres)", f"{max_acres:,.1f}"),
        ("Number of National Forests Represented", f"{n_forests:,}"),
        ("Number of States Represented", f"{n_states:,}"),
        ("Number of USFS Regions Represented", f"{n_regions:,}"),
    ]

    return pd.DataFrame(data, columns=["Metric", "Value"])


def build_table2_by_region(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    total_acres = gdf["ACRES"].sum(skipna=True)

    grouped = (
        gdf
        .groupby("REGION", dropna=False)
        .agg(
            Num_Polygons=("NAME", "count"),
            Total_Acres=("ACRES", "sum"),
            Mean_Acres=("ACRES", "mean"),
        )
        .reset_index()
    )

    grouped["Percent_of_National_Acres"] = (
        grouped["Total_Acres"] / total_acres * 100
    )

    # Sort by acres descending for readability
    grouped = grouped.sort_values("Total_Acres", ascending=False)

    # Format REGION as string (handles NaN)
    grouped["REGION"] = grouped["REGION"].astype(str)

    # Pretty rounding
    grouped["Total_Acres"] = grouped["Total_Acres"].round(1)
    grouped["Mean_Acres"] = grouped["Mean_Acres"].round(1)
    grouped["Percent_of_National_Acres"] = grouped["Percent_of_National_Acres"].round(1)

    return grouped.rename(
        columns={
            "REGION": "USFS Region",
            "Num_Polygons": "Number of Polygons",
            "Total_Acres": "Total Acres",
            "Mean_Acres": "Mean Acres",
            "Percent_of_National_Acres": "% of National Total",
        }
    )


def build_table3_by_state(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    total_acres = gdf["ACRES"].sum(skipna=True)

    # Largest polygon per state
    max_acres_by_state = (
        gdf.groupby("STATE")["ACRES"].max().rename("Largest_Acres")
    )

    grouped = (
        gdf
        .groupby("STATE", dropna=False)
        .agg(
            Num_Polygons=("NAME", "count"),
            Total_Acres=("ACRES", "sum"),
            Mean_Acres=("ACRES", "mean"),
        )
        .reset_index()
        .merge(max_acres_by_state.reset_index(), on="STATE", how="left")
    )

    grouped["Percent_of_National_Acres"] = (
        grouped["Total_Acres"] / total_acres * 100
    )

    grouped = grouped.sort_values("Total_Acres", ascending=False)

    # Formatting
    for col in ["Total_Acres", "Mean_Acres", "Largest_Acres"]:
        grouped[col] = grouped[col].round(1)

    grouped["Percent_of_National_Acres"] = grouped["Percent_of_National_Acres"].round(1)

    return grouped.rename(
        columns={
            "STATE": "State",
            "Num_Polygons": "Number of Polygons",
            "Total_Acres": "Total Acres",
            "Mean_Acres": "Mean Acres",
            "Largest_Acres": "Largest Roadless Area (acres)",
            "Percent_of_National_Acres": "% of National Total",
        }
    )


def build_table4_by_forest(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    grouped = (
        gdf
        .groupby(["FOREST", "STATE"], dropna=False)
        .agg(
            Num_Polygons=("NAME", "count"),
            Total_Acres=("ACRES", "sum"),
            Mean_Acres=("ACRES", "mean"),
        )
        .reset_index()
    )

    grouped = grouped.sort_values("Total_Acres", ascending=False)

    # Formatting
    grouped["Total_Acres"] = grouped["Total_Acres"].round(1)
    grouped["Mean_Acres"] = grouped["Mean_Acres"].round(1)

    return grouped.rename(
        columns={
            "FOREST": "National Forest",
            "STATE": "State(s)",
            "Num_Polygons": "Number of Polygons",
            "Total_Acres": "Total Acres",
            "Mean_Acres": "Mean Acres",
        }
    )


def build_table5_by_category(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    total_acres = gdf["ACRES"].sum(skipna=True)

    grouped = (
        gdf
        .groupby("CATEGORY", dropna=False)
        .agg(
            Num_Polygons=("NAME", "count"),
            Total_Acres=("ACRES", "sum"),
            Mean_Acres=("ACRES", "mean"),
        )
        .reset_index()
    )

    grouped["Percent_of_National_Acres"] = (
        grouped["Total_Acres"] / total_acres * 100
    )

    grouped = grouped.sort_values("Total_Acres", ascending=False)

    # Formatting
    grouped["Total_Acres"] = grouped["Total_Acres"].round(1)
    grouped["Mean_Acres"] = grouped["Mean_Acres"].round(1)
    grouped["Percent_of_National_Acres"] = grouped["Percent_of_National_Acres"].round(1)

    return grouped.rename(
        columns={
            "CATEGORY": "Category",
            "Num_Polygons": "Number of Polygons",
            "Total_Acres": "Total Acres",
            "Mean_Acres": "Mean Acres",
            "Percent_of_National_Acres": "% of National Total",
        }
    )


def build_table6_missing_values(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    fields = [
        "REGION", "FOREST", "STATE", "NAME",
        "CATEGORY", "ACRES", "SHAPE_AREA", "SHAPE_LEN", "geometry"
    ]

    total_rows = len(gdf)
    rows = []
    for field in fields:
        n_missing = gdf[field].isna().sum()
        pct_missing = n_missing / total_rows * 100 if total_rows > 0 else np.nan
        rows.append((field, n_missing, pct_missing))

    df = pd.DataFrame(rows, columns=["Field", "Missing Values", "% Missing"])
    df["% Missing"] = df["% Missing"].round(2)
    return df


def build_table7_geometry_stats(gdf: gpd.GeoDataFrame) -> pd.DataFrame:
    stats = {}
    for col in ["SHAPE_AREA", "SHAPE_LEN"]:
        desc = gdf[col].describe()
        stats[col] = {
            "Min": desc["min"],
            "Max": desc["max"],
            "Mean": desc["mean"],
            "Median": desc["50%"],
            "Std Dev": desc["std"],
        }

    table = pd.DataFrame(stats).T.reset_index().rename(columns={"index": "Metric"})
    # Round values reasonably
    for c in ["Min", "Max", "Mean", "Median", "Std Dev"]:
        table[c] = table[c].round(3)

    return table.rename(columns={"Metric": "Geometry Field"})


# ------------------ EXPORT HELPERS ------------------
def export_df_to_all_formats(
    df: pd.DataFrame,
    base_name: str,
    output_dir: Path,
) -> None:
    """
    Export a DataFrame to CSV, Markdown, and LaTeX.
    Files will be named:
      base_name.csv, base_name.md, base_name.tex
    """
    # CSV
    csv_path = output_dir / f"{base_name}.csv"
    df.to_csv(csv_path, index=False)

    # Markdown (requires tabulate installed)
    md_path = output_dir / f"{base_name}.md"
    try:
        md_str = df.to_markdown(index=False)
    except Exception:
        # Fallback simple markdown if to_markdown not available
        md_str = df.to_csv(index=False, sep="|")
    md_path.write_text(md_str, encoding="utf-8")

    # LaTeX
    tex_path = output_dir / f"{base_name}.tex"
    tex_str = df.to_latex(index=False, escape=True)
    tex_path.write_text(tex_str, encoding="utf-8")

def add_df_to_word_doc(doc: DocxDocument, df: pd.DataFrame, title: str) -> None:
    """
    Append a title and a table (from df) to a python-docx Document.
    """
    doc.add_heading(title, level=2)

    # Create table with header row
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            row_cells[j].text = "" if pd.isna(row[col]) else str(row[col])

    doc.add_paragraph()  # spacing after table
    doc.add_paragraph()  # spacing after table


def export_all_tables_to_word(tables: dict[str, pd.DataFrame], output_path: Path) -> None:
    """
    tables: dict mapping table label -> DataFrame
    """
    doc = Document()
    doc.add_heading("Roadless Area Inventory – Summary Tables", level=1)

    for label, df in tables.items():
        add_df_to_word_doc(doc, df, label)

    doc.save(str(output_path))


# ------------------------ MAIN ----------------------
def main():
    print("Loading roadless dataset...")
    gdf = load_roadless(ROADLESS_PATH, ROADLESS_LAYER)

    print("Building tables...")
    t1 = build_table1_national_summary(gdf)
    t2 = build_table2_by_region(gdf)
    t3 = build_table3_by_state(gdf)
    t4 = build_table4_by_forest(gdf)
    t5 = build_table5_by_category(gdf)
    t6 = build_table6_missing_values(gdf)
    t7 = build_table7_geometry_stats(gdf)

    tables = {
        "Table 1. National Roadless Area Summary": t1,
        "Table 2. Roadless Areas by USFS Region": t2,
        "Table 3. Roadless Areas by State": t3,
        "Table 4. Roadless Areas by National Forest": t4,
        "Table 5. Roadless Areas by Category": t5,
        "Table 6. Attribute Completeness Summary": t6,
        "Table 7. Geometry Statistics": t7,
    }

    print("Exporting individual tables to CSV / Markdown / LaTeX...")
    export_df_to_all_formats(t1, "table1_national_summary", OUTPUT_DIR)
    export_df_to_all_formats(t2, "table2_by_region", OUTPUT_DIR)
    export_df_to_all_formats(t3, "table3_by_state", OUTPUT_DIR)
    export_df_to_all_formats(t4, "table4_by_forest", OUTPUT_DIR)
    export_df_to_all_formats(t5, "table5_by_category", OUTPUT_DIR)
    export_df_to_all_formats(t6, "table6_missing_values", OUTPUT_DIR)
    export_df_to_all_formats(t7, "table7_geometry_stats", OUTPUT_DIR)

    print("Exporting all tables into a single Word document...")
    word_path = OUTPUT_DIR / "roadless_inventory_tables.docx"
    export_all_tables_to_word(tables, word_path)

    print("Done.")
    print(f"Outputs written to: {OUTPUT_DIR.resolve()}")


if __name__ == "__main__":
    main()
