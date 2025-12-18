"""
Roadless Areas – Forest to Faucets (F2F)
Top 20 Roadless Areas by Drinking-Water Importance

Outputs:
- CSV tables
- LaTeX tables
- Single Word document with all tables

No Markdown, no optional dependencies.
"""

# ---------------------- IMPORTS ----------------------
from pathlib import Path
import pandas as pd
import geopandas as gpd
import numpy as np
from docx import Document
from docx.document import Document as DocxDocument

from nfw_project import config_paths as cfg


# ---------------------- CONFIG ----------------------
OUTPUT_DIR = Path(cfg.TABLES_OUTPUT_DIR) / "roadless_f2f_tables"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

ROADLESS_LAYER = "roadless_area"
F2F_LAYER = "F2F_RA"


# ------------------ EXPORT HELPERS ------------------
def export_df_to_all_formats(
    df: pd.DataFrame,
    base_name: str,
    output_dir: Path,
) -> None:
    """
    Export a DataFrame to CSV, Markdown, and LaTeX.
    Files will be named:
      base_name.csv, base_name.md, base_name.tex
    """
    # CSV
    csv_path = output_dir / f"{base_name}.csv"
    df.to_csv(csv_path, index=False)

    # Markdown (requires tabulate installed)
    md_path = output_dir / f"{base_name}.md"
    try:
        md_str = df.to_markdown(index=False)
    except Exception:
        # Fallback simple markdown if to_markdown not available
        md_str = df.to_csv(index=False, sep="|")
    md_path.write_text(md_str, encoding="utf-8")

    # LaTeX
    tex_path = output_dir / f"{base_name}.tex"
    tex_str = df.to_latex(index=False, escape=True)
    tex_path.write_text(tex_str, encoding="utf-8")

def add_df_to_word_doc(doc: DocxDocument, df: pd.DataFrame, title: str) -> None:
    """
    Append a title and a table (from df) to a python-docx Document.
    """
    doc.add_heading(title, level=2)

    # Create table with header row
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            row_cells[j].text = "" if pd.isna(row[col]) else str(row[col])

    doc.add_paragraph()  # spacing after table
    doc.add_paragraph()  # spacing after table


def export_all_tables_to_word(tables: dict[str, pd.DataFrame], output_path: Path) -> None:
    """
    tables: dict mapping table label -> DataFrame
    """
    doc = Document()
    doc.add_heading("Roadless Areas – Drinking Water Importance (Forest to Faucets)", level=1)

    for label, df in tables.items():
        add_df_to_word_doc(doc, df, label)

    doc.save(str(output_path))


# ------------------ DATA LOADING ------------------
def load_data():
    roadless = gpd.read_file(cfg.USFS_GPKG, layer=ROADLESS_LAYER)
    f2f = gpd.read_file(cfg.ROADLESS_ANALYSIS_GPKG, layer=F2F_LAYER)

    if roadless.crs != f2f.crs:
        f2f = f2f.to_crs("EPSG:5070")

    if "RA_ID" not in roadless.columns:
        roadless = roadless.reset_index(drop=True)
        roadless["RA_ID"] = roadless.index + 1

    return roadless, f2f


# ------------------ ANALYSIS ------------------
def coerce_f2f_numeric(F2F: gpd.GeoDataFrame) -> gpd.GeoDataFrame:
    # ensure numeric
    num_cols = [
        "ACRES", "Domestic", "Industrial", "Irrigation", "Livestock", "Mining",
        "Thermo", "Public_sup", "Aquacultur", "Total_SW",
        "Ps_del_dom", "Domestic_GW", "Industri_GW", "Irrigati_GW",
        "Livestoc_GW", "Mining_GW", "Thermo_GW", "Public_sup_GW",
        "Aquacult_GW", "Total_GW"
    ]
    for c in num_cols:
        if c in F2F.columns:
            F2F[c] = pd.to_numeric(F2F[c], errors="coerce")
    return F2F


def build_f2f_huc_summary(F2F: gpd.GeoDataFrame) -> pd.DataFrame:
    # HUC8-level inventory (these are *already* intersecting RAs)
    f2f_huc_summary = (
        F2F
        .groupby(["HUC_8", "HU_8_STATE"], dropna=False)
        .agg(
            HUC8_Acres=("ACRES", "first"),     # F2F polygon acres
            Total_SW=("Total_SW", "first"),
            Total_GW=("Total_GW", "first"),
            Public_sup_SW=("Public_sup", "first"),
            Public_sup_GW=("Public_sup_GW", "first"),
        )
        .reset_index()
    )
    return f2f_huc_summary


def ensure_roadless_ids(roadless_5070: gpd.GeoDataFrame) -> gpd.GeoDataFrame:
    # keep a stable ID for each roadless polygon
    if "RA_ID" not in roadless_5070.columns:
        roadless_5070 = roadless_5070.reset_index(drop=True)
        roadless_5070["RA_ID"] = roadless_5070.index + 1  # simple integer ID
    return roadless_5070


def build_ra_huc_overlay(
    roadless_5070: gpd.GeoDataFrame,
    F2F: gpd.GeoDataFrame
) -> gpd.GeoDataFrame:
    # Make sure both in same CRS (use 5070 / your analysis CRS)
    if roadless_5070.crs is None:
        raise ValueError("roadless_5070.crs is None; set CRS before overlay.")
    if F2F.crs is None:
        raise ValueError("F2F.crs is None; set CRS before overlay.")

    # ensure everything is EPSG:5070 for area calcs
    if str(roadless_5070.crs).lower() != "epsg:5070":
        roadless_5070 = roadless_5070.to_crs("EPSG:5070")
    if str(F2F.crs).lower() != "epsg:5070":
        F2F = F2F.to_crs("EPSG:5070")

    # Spatial overlay – splits geometries into RA–HUC pieces
    ra_huc = gpd.overlay(
        roadless_5070[["RA_ID", "REGION", "FOREST", "STATE", "NAME", "geometry"]],
        F2F[[
            "HUC_8", "HU_8_STATE", "ACRES",
            "Total_SW", "Total_GW", "Public_sup", "Public_sup_GW",
            "geometry"
        ]],
        how="intersection"
    )

    # Area of the intersection in acres (assuming CRS is in meters)
    ra_huc["INT_AREA_M2"] = ra_huc.geometry.area
    ra_huc["INT_ACRES"] = ra_huc["INT_AREA_M2"] / 4046.8564224

    # Fraction of the HUC8 inside the roadless area
    # (Assumes F2F['ACRES'] is HUC8 polygon area in acres)
    ra_huc["FRAC_HUC_IN_RA"] = ra_huc["INT_ACRES"] / ra_huc["ACRES"]

    # Area-weighted metrics
    water_fields = ["Total_SW", "Total_GW", "Public_sup", "Public_sup_GW"]
    for fld in water_fields:
        if fld in ra_huc.columns:
            ra_huc[f"{fld}_w"] = ra_huc[fld] * ra_huc["FRAC_HUC_IN_RA"]

    return ra_huc


def build_ra_water_summary(ra_huc: gpd.GeoDataFrame) -> pd.DataFrame:
    ra_water = (
        ra_huc
        .groupby(["RA_ID", "REGION", "FOREST", "STATE", "NAME"], dropna=False)
        .agg(
            RA_Acres=("INT_ACRES", "sum"),   # acres of RA that lie in F2F HUCs
            HUC8_Count=("HUC_8", "nunique"),
            Total_SW_w=("Total_SW_w", "sum"),
            Total_GW_w=("Total_GW_w", "sum"),
            Public_sup_w=("Public_sup_w", "sum"),
            Public_sup_GW_w=("Public_sup_GW_w", "sum"),
        )
        .reset_index()
    )

    # total public supply
    ra_water["Public_sup_Total_w"] = (
        ra_water["Public_sup_w"].fillna(0) +
        ra_water["Public_sup_GW_w"].fillna(0)
    )

    return ra_water


    top20_area_weighted_public_supply = (
        ra_water
        .sort_values("Public_sup_Total_w", ascending=False)
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Public_sup_Total_w"
        ]]
    )
    top20_area_weighted_public_supply["RA_Acres"] = top20_area_weighted_public_supply["RA_Acres"].round(1)
    top20_area_weighted_public_supply["Public_sup_Total_w"] = top20_area_weighted_public_supply["Public_sup_Total_w"].round(2)


    # 2) Public supply per acre (small-but-critical headwaters)
    ra_water["Public_sup_per_RA_acre"] = (
        ra_water["Public_sup_Total_w"] / ra_water["RA_Acres"].replace(0, np.nan)
    )

    top20_public_supply_per_acre = (
        ra_water
        .sort_values("Public_sup_per_RA_acre", ascending=False)
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Public_sup_per_RA_acre", "Public_sup_Total_w"
        ]]
    )
    top20_public_supply_per_acre["RA_Acres"] = top20_public_supply_per_acre["RA_Acres"].round(1)
    top20_public_supply_per_acre["Public_sup_per_RA_acre"] = top20_public_supply_per_acre["Public_sup_per_RA_acre"].round(4)
    top20_public_supply_per_acre["Public_sup_Total_w"] = top20_public_supply_per_acre["Public_sup_Total_w"].round(2)


    # 3) Total withdrawals (surface + groundwater)
    ra_water["Total_withdrawals_w"] = (
        ra_water["Total_SW_w"].fillna(0) + ra_water["Total_GW_w"].fillna(0)
    )

    top20_total_withdrawals = (
        ra_water
        .sort_values("Total_withdrawals_w", ascending=False)
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Total_withdrawals_w", "Total_SW_w", "Total_GW_w"
        ]]
    )
    top20_total_withdrawals["RA_Acres"] = top20_total_withdrawals["RA_Acres"].round(1)
    for c in ["Total_withdrawals_w", "Total_SW_w", "Total_GW_w"]:
        top20_total_withdrawals[c] = top20_total_withdrawals[c].round(2)


    # 4) Multi-HUC exposure (watershed reach)
    top20_by_huc8_count = (
        ra_water
        .sort_values(["HUC8_Count", "Public_sup_Total_w"], ascending=[False, False])
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Public_sup_Total_w"
        ]]
    )
    top20_by_huc8_count["RA_Acres"] = top20_by_huc8_count["RA_Acres"].round(1)
    top20_by_huc8_count["Public_sup_Total_w"] = top20_by_huc8_count["Public_sup_Total_w"].round(2)


    # 5) Groundwater reliance (share of public supply that is GW)
    ra_water["Public_sup_GW_share"] = (
        ra_water["Public_sup_GW_w"].fillna(0) / ra_water["Public_sup_Total_w"].replace(0, np.nan)
    )

    top20_public_supply_gw_share = (
        ra_water
        .sort_values("Public_sup_GW_share", ascending=False)
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Public_sup_GW_share", "Public_sup_Total_w"
        ]]
    )
    top20_public_supply_gw_share["RA_Acres"] = top20_public_supply_gw_share["RA_Acres"].round(1)
    top20_public_supply_gw_share["Public_sup_GW_share"] = top20_public_supply_gw_share["Public_sup_GW_share"].round(3)
    top20_public_supply_gw_share["Public_sup_Total_w"] = top20_public_supply_gw_share["Public_sup_Total_w"].round(2)


# ------------------------ MAIN ----------------------
def main():
    print("Loading datasets...")
    roadless_5070, F2F = load_data()  # expects (roadless, f2f)

    print("Coercing F2F numeric fields...")
    F2F = coerce_f2f_numeric(F2F)

    print("Building HUC8-level F2F summary...")
    f2f_huc_summary = build_f2f_huc_summary(F2F)
    n_huc = f2f_huc_summary["HUC_8"].nunique()
    print(f"{n_huc} unique HUC8 watersheds intersect at least one roadless area.")

    print("Ensuring Roadless IDs...")
    roadless_5070 = ensure_roadless_ids(roadless_5070)

    print("Building RA–HUC overlay...")
    ra_huc = build_ra_huc_overlay(roadless_5070, F2F)

    print("Aggregating to roadless areas...")
    ra_water = build_ra_water_summary(ra_huc)

    # ------------------ TOP-20 METRIC TABLES ------------------

    # 1) Headline metric (rename: not "drinking-water importance")
    top20_area_weighted_public_supply = (
        ra_water
        .sort_values("Public_sup_Total_w", ascending=False)
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Public_sup_Total_w"
        ]]
    )
    top20_area_weighted_public_supply["RA_Acres"] = top20_area_weighted_public_supply["RA_Acres"].round(1)
    top20_area_weighted_public_supply["Public_sup_Total_w"] = top20_area_weighted_public_supply["Public_sup_Total_w"].round(2)

    # 2) Public supply per acre (small-but-critical headwaters)
    ra_water["Public_sup_per_RA_acre"] = (
        ra_water["Public_sup_Total_w"] / ra_water["RA_Acres"].replace(0, np.nan)
    )

    top20_public_supply_per_acre = (
        ra_water
        .sort_values("Public_sup_per_RA_acre", ascending=False)
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Public_sup_per_RA_acre", "Public_sup_Total_w"
        ]]
    )
    top20_public_supply_per_acre["RA_Acres"] = top20_public_supply_per_acre["RA_Acres"].round(1)
    top20_public_supply_per_acre["Public_sup_per_RA_acre"] = top20_public_supply_per_acre["Public_sup_per_RA_acre"].round(4)
    top20_public_supply_per_acre["Public_sup_Total_w"] = top20_public_supply_per_acre["Public_sup_Total_w"].round(2)

    # 3) Total withdrawals (surface + groundwater)
    ra_water["Total_withdrawals_w"] = (
        ra_water["Total_SW_w"].fillna(0) + ra_water["Total_GW_w"].fillna(0)
    )

    top20_total_withdrawals = (
        ra_water
        .sort_values("Total_withdrawals_w", ascending=False)
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Total_withdrawals_w", "Total_SW_w", "Total_GW_w"
        ]]
    )
    top20_total_withdrawals["RA_Acres"] = top20_total_withdrawals["RA_Acres"].round(1)
    for c in ["Total_withdrawals_w", "Total_SW_w", "Total_GW_w"]:
        top20_total_withdrawals[c] = top20_total_withdrawals[c].round(2)

    # 4) Multi-HUC exposure (watershed reach)
    top20_by_huc8_count = (
        ra_water
        .sort_values(["HUC8_Count", "Public_sup_Total_w"], ascending=[False, False])
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Public_sup_Total_w"
        ]]
    )
    top20_by_huc8_count["RA_Acres"] = top20_by_huc8_count["RA_Acres"].round(1)
    top20_by_huc8_count["Public_sup_Total_w"] = top20_by_huc8_count["Public_sup_Total_w"].round(2)

    # 5) Groundwater reliance (share of public supply that is GW)
    ra_water["Public_sup_GW_share"] = (
        ra_water["Public_sup_GW_w"].fillna(0) / ra_water["Public_sup_Total_w"].replace(0, np.nan)
    )

    top20_public_supply_gw_share = (
        ra_water
        .sort_values("Public_sup_GW_share", ascending=False)
        .head(20)
        .assign(Rank=lambda df: range(1, len(df) + 1))
        [[
            "Rank", "NAME", "FOREST", "STATE", "REGION",
            "RA_Acres", "HUC8_Count", "Public_sup_GW_share", "Public_sup_Total_w"
        ]]
    )
    top20_public_supply_gw_share["RA_Acres"] = top20_public_supply_gw_share["RA_Acres"].round(1)
    top20_public_supply_gw_share["Public_sup_GW_share"] = top20_public_supply_gw_share["Public_sup_GW_share"].round(3)
    top20_public_supply_gw_share["Public_sup_Total_w"] = top20_public_supply_gw_share["Public_sup_Total_w"].round(2)

    # ------------------ EXPORT (CSV + LaTeX) ------------------
    print("Exporting Top-20 tables to CSV / LaTeX...")

    export_df_to_all_formats(top20_area_weighted_public_supply,
                             "top20_area_weighted_public_supply_f2f",
                             OUTPUT_DIR)

    export_df_to_all_formats(top20_public_supply_per_acre,
                             "top20_public_supply_per_acre_f2f",
                             OUTPUT_DIR)

    export_df_to_all_formats(top20_total_withdrawals,
                             "top20_area_weighted_total_withdrawals_f2f",
                             OUTPUT_DIR)

    export_df_to_all_formats(top20_by_huc8_count,
                             "top20_multi_huc_exposure_f2f",
                             OUTPUT_DIR)

    export_df_to_all_formats(top20_public_supply_gw_share,
                             "top20_public_supply_gw_share_f2f",
                             OUTPUT_DIR)

    # ------------------ WORD DOC (ALL TOP-20 TABLES) ------------------
    print("Exporting all Top-20 tables to Word...")

    tables_word = {
        "Top 20 Roadless Areas by Area-Weighted Public Supply (F2F HUC8 Proxy)":
            top20_area_weighted_public_supply,

        "Top 20 Roadless Areas by Area-Weighted Public Supply per Roadless Acre (F2F)":
            top20_public_supply_per_acre,

        "Top 20 Roadless Areas by Area-Weighted Total Withdrawals (Surface + Groundwater)":
            top20_total_withdrawals,

        "Top 20 Roadless Areas by Multi-HUC8 Exposure (Number of Intersecting Watersheds)":
            top20_by_huc8_count,

        "Top 20 Roadless Areas by Groundwater Reliance (Share of Public Supply from GW)":
            top20_public_supply_gw_share,
    }

    export_all_tables_to_word(
        tables_word,
        OUTPUT_DIR / "roadless_f2f_top20_tables.docx"
    )

    print("Done.")
    print(f"Outputs written to: {OUTPUT_DIR.resolve()}")



if __name__ == "__main__":
    main()
